"""
word_import_service.py
Parses pasted text blocks and real .docx files into resident dicts.
Supports paragraph blocks (Key: Value / Key - Value) and table extraction.
Safe to import and test standalone.
"""

"""
word_import_service.py
Parses pasted text blocks and real .docx files into resident dicts.
Supports paragraph blocks (Key: Value / Key - Value) and table extraction.
Safe to import and test standalone.
"""

import io
import re
import docx
from typing import Any

# Canonical field aliases (all lowercase for matching).
# Kept generous on purpose — Gram Panchayat paperwork uses many phrasings
# for the same field ("Property No", "Ward No", "Mobile Number", "Tax Amt"...).
FIELD_MAP: dict[str, list[str]] = {
    "name": [
        "name", "taxpayer", "taxpayer name", "resident name", "owner name",
        "owner", "full name", "नाव", "करदात्याचे नाव", "मालकाचे नाव",
    ],
    "property_id": [
        "property id", "property", "id", "prop id", "property_id",
        "property no", "property number", "property no.", "prop no",
        "house no", "house number", "house no.",
        "मालमत्ता क्रमांक", "घर क्र", "घर क्रमांक",
    ],
    "ward": [
        "ward", "ward no", "ward no.", "ward number",
        "प्रभाग", "वॉर्ड", "वॉर्ड क्र",
    ],
    "phone": [
        "phone", "mobile", "mob", "mobile no", "mobile no.", "mobile number",
        "phone no", "phone no.", "phone number", "contact", "contact no",
        "contact no.", "contact number", "cell", "cell no",
        "मोबाईल", "फोन", "संपर्क",
    ],
    "address": [
        "address", "addr", "location", "पत्ता", "ठिकाण",
    ],
    "base_amount": [
        "amount", "tax amount", "base amount", "tax", "base_amount",
        "tax amt", "tax amt.", "amt", "amount due", "property tax",
        "total tax", "tax due",
        "रक्कम", "कर रक्कम", "एकूण रक्कम", "मालमत्ता कर",
    ],
}


def _normalize_key(raw_key: str) -> str:
    """
    Normalizes a parsed header/key before alias lookup so minor
    punctuation/spacing differences ('Property No.', 'Property No',
    'property  no') all resolve the same way.
    """
    k = raw_key.strip().lower()
    k = re.sub(r"[.\-_]+", " ", k)   # dots/hyphens/underscores -> space
    k = re.sub(r"\s+", " ", k).strip()
    return k


def _match_field(raw_key: str) -> str | None:
    """Looks up a normalized key against FIELD_MAP, returns the canonical field or None."""
    norm = _normalize_key(raw_key)
    for field, aliases in FIELD_MAP.items():
        for alias in aliases:
            if norm == _normalize_key(alias):
                return field
    return None


def _clean_amount(raw: str) -> float:
    """
    Extract a numeric amount from strings like:
      '₹2,500'  '2,500.50'  'Rs. 2500/-'  'Rs.2,500'  '2,50,000'  'INR 1200'
    Strategy: drop known currency words/symbols and trailing '/-' first,
    THEN strip remaining non-digit/comma/dot characters, THEN take the
    numeric run — this avoids stray dots in abbreviations like 'Rs.'
    being misread as a decimal point.
    Returns 0.0 on failure.
    """
    if not raw:
        return 0.0

    s = raw.strip()
    # Remove common currency markers/words (case-insensitive) and trailing "/-"
    s = re.sub(r"(?i)\b(rs|inr|rupees?)\b\.?", "", s)
    s = s.replace("₹", "")
    s = re.sub(r"/-\s*$", "", s)
    s = s.strip()

    # Now pull out the first plausible number: digits with optional
    # comma grouping and at most one decimal point.
    m = re.search(r"\d[\d,]*(?:\.\d+)?", s)
    if not m:
        return 0.0

    cleaned = m.group(0).replace(",", "")
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def parse_text(raw_text: str) -> list[dict[str, Any]]:
    """
    Split raw pasted text into blocks (separated by one or more blank lines).
    Each block is parsed line-by-line using 'Key: Value' or 'Key - Value' format.
    Returns only records that have at minimum name + property_id.
    Missing fields are stored as empty strings / 0.0.
    """
    blocks = re.split(r"\n\s*\n", raw_text.strip())
    results: list[dict[str, Any]] = []

    for block_idx, block in enumerate(blocks):
        record: dict[str, Any] = {
            "name":        "",
            "property_id": "",
            "ward":        "",
            "phone":       "",
            "address":     "",
            "base_amount": 0.0,
            "_block_index": block_idx,
            "_raw_block":   block.strip(),
        }

        for line in block.strip().splitlines():
            line = line.strip()
            if not line:
                continue

            # Match   Key: Value   or   Key - Value
            m = re.match(r"^(.+?)\s*[:\-]\s*(.+)$", line, re.IGNORECASE)
            if not m:
                continue

            raw_key = m.group(1).strip()
            value   = m.group(2).strip()

            matched_field = _match_field(raw_key)

            if matched_field is None:
                continue  # unrecognised key — skip silently

            if matched_field == "base_amount":
                record[matched_field] = _clean_amount(value)
            else:
                record[matched_field] = value

        # Only keep records that have the two required fields
        if record["name"] and record["property_id"]:
            results.append(record)

    return results


def parse_docx_bytes(docx_bytes: bytes) -> list[dict[str, Any]]:
    """Parse .docx binary content from file upload."""
    doc = docx.Document(io.BytesIO(docx_bytes))
    return _parse_docx_doc(doc)


def parse_docx_file(file_path: str) -> list[dict[str, Any]]:
    """Parse .docx directly from filesystem path."""
    doc = docx.Document(file_path)
    return _parse_docx_doc(doc)


def _parse_docx_doc(doc: docx.Document) -> list[dict[str, Any]]:
    """
    Extracts records from both paragraph text blocks and tables inside a Word .docx document.
    """
    records: list[dict[str, Any]] = []
    seen_ids = set()

    # 1. Paragraph blocks
    paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
    if paragraphs_text.strip():
        parsed_paras = parse_text(paragraphs_text)
        for r in parsed_paras:
            pid = r["property_id"].strip().lower()
            if pid and pid not in seen_ids:
                seen_ids.add(pid)
                records.append(r)

    # 2. Tables inside the docx
    for table in doc.tables:
        if not table.rows:
            continue

        # Check for column headers in first row
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        col_map: dict[int, str] = {}
        for col_idx, header in enumerate(header_cells):
            field = _match_field(header)
            if field:
                col_map[col_idx] = field

        if "name" in col_map.values() and "property_id" in col_map.values():
            for row_idx, row in enumerate(table.rows[1:], start=1):
                rec: dict[str, Any] = {
                    "name": "", "property_id": "", "ward": "", "phone": "",
                    "address": "", "base_amount": 0.0,
                    "_block_index": len(records),
                    "_raw_block": f"Table Row {row_idx}"
                }
                for col_idx, cell in enumerate(row.cells):
                    field = col_map.get(col_idx)
                    if field:
                        val = cell.text.strip()
                        if field == "base_amount":
                            rec[field] = _clean_amount(val)
                        else:
                            rec[field] = val

                pid = rec["property_id"].strip().lower()
                if rec["name"] and rec["property_id"] and pid not in seen_ids:
                    seen_ids.add(pid)
                    records.append(rec)
        else:
            # Try parsing key-value blocks in table text
            table_text = ""
            for row in table.rows:
                table_text += "\n".join(cell.text.strip() for cell in row.cells if cell.text.strip()) + "\n"
            if table_text.strip():
                t_recs = parse_text(table_text)
                for tr in t_recs:
                    pid = tr["property_id"].strip().lower()
                    if pid and pid not in seen_ids:
                        seen_ids.add(pid)
                        records.append(tr)

    return records


def check_duplicates(
    parsed_list: list[dict],
    existing_property_ids: list[str],
) -> tuple[list[dict], list[dict]]:
    """
    Separate parsed records into:
      - new_records:       property_id not in existing set → safe to insert
      - duplicate_records: property_id already exists → needs admin review
    """
    existing_set = set(pid.strip().lower() for pid in existing_property_ids)
    new_records: list[dict] = []
    duplicate_records: list[dict] = []

    for rec in parsed_list:
        if rec["property_id"].strip().lower() in existing_set:
            duplicate_records.append(rec)
        else:
            new_records.append(rec)

    return new_records, duplicate_records


# ── DB-facing functions ───────────────────────────────────────────────────────

def confirm_import(new_records: list[dict]) -> dict:
    """Insert confirmed-new records into the DB."""
    from database.db import execute_write_many
    statements: list[tuple] = []
    errors: list[str] = []

    for rec in new_records:
        if not rec.get("name") or not rec.get("property_id"):
            errors.append(f"Skipped incomplete record: {rec.get('_raw_block', '')[:40]}")
            continue
        statements.append((
            """INSERT INTO residents
               (name, property_id, ward, phone, address, base_amount, payment_status)
               VALUES (?, ?, ?, ?, ?, ?, 'unpaid')""",
            (
                rec["name"],
                rec["property_id"],
                rec.get("ward", ""),
                rec.get("phone", ""),
                rec.get("address", ""),
                rec.get("base_amount", 0.0),
            )
        ))

    if statements:
        execute_write_many(statements)

    return {"inserted": len(statements), "errors": errors}


def merge_resident(existing_id: int, new_data: dict) -> dict:
    """Staff-approved overwrite of an existing resident's mutable fields."""
    from database.db import execute_write
    execute_write(
        """UPDATE residents
           SET name=?, ward=?, phone=?, address=?, base_amount=?
           WHERE id=?""",
        (
            new_data.get("name", ""),
            new_data.get("ward", ""),
            new_data.get("phone", ""),
            new_data.get("address", ""),
            new_data.get("base_amount", 0.0),
            existing_id,
        )
    )
    return {"merged": True, "id": existing_id}
